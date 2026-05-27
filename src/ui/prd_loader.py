"""Extract plain-text PRD content from an uploaded file.

Streamlit's `st.file_uploader` hands us a `BytesIO`-like object plus a
filename. We dispatch by extension:

  .pdf            → pypdf, page-by-page text extraction
  .docx           → python-docx, paragraph join
  .md / .txt      → utf-8 decode (best-effort: latin-1 fallback)
  anything else   → raise UnsupportedFileType

We don't try to handle legacy `.doc` (Microsoft binary format) — no
pure-Python library exists that's worth shipping. The UI tells the
user to "export to .docx or .pdf" instead.

Failure modes:
  - Empty extraction (e.g. scanned PDF with no embedded text) raises
    EmptyExtractionError so the UI can suggest OCR.
  - Files over `MAX_BYTES` raise FileTooLargeError. The cap protects
    the demo from someone uploading a 50 MB design doc.
"""

from __future__ import annotations

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


# Cap upload size to keep the demo responsive and to bound LLM context.
# A typical PRD is 1–10 KB of text; 2 MB allows for image-heavy PDFs
# where the embedded text is still small.
MAX_BYTES = 2 * 1024 * 1024


class PRDLoaderError(Exception):
    """Base class for all extraction errors."""


class UnsupportedFileType(PRDLoaderError):
    pass


class FileTooLargeError(PRDLoaderError):
    pass


class EmptyExtractionError(PRDLoaderError):
    """Parsed the file but got nothing useful (e.g. scanned PDF with no text layer)."""


def extract_text(filename: str, data: bytes) -> str:
    """Extract plain-text PRD content from `data`, dispatching by extension."""
    if len(data) > MAX_BYTES:
        raise FileTooLargeError(
            f"文件超过 {MAX_BYTES // (1024 * 1024)} MB 上限，请压缩后再试"
        )

    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        text = _extract_pdf(data)
    elif ext == ".docx":
        text = _extract_docx(data)
    elif ext in {".md", ".markdown", ".txt"}:
        text = _extract_plain(data)
    elif ext == ".doc":
        raise UnsupportedFileType(
            "旧版 .doc 文件无法解析，请在 Word 里另存为 .docx 后再上传"
        )
    else:
        raise UnsupportedFileType(
            f"暂不支持 {ext or '该格式'} 文件，请用 PDF / Word(.docx) / Markdown / TXT"
        )

    text = text.strip()
    if not text:
        raise EmptyExtractionError(
            "未能从文件中提取出文字。如果这是扫描版 PDF，请先 OCR 后再上传。"
        )
    return text


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader  # local import — keeps top-level fast

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001
        raise PRDLoaderError(f"PDF 解析失败：{e}") from e

    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            parts.append(page.extract_text() or "")
        except Exception as e:  # noqa: BLE001
            logger.warning("PDF page %d extraction failed: %s", i, e)
    return "\n\n".join(p for p in parts if p.strip())


def _extract_docx(data: bytes) -> str:
    # `python-docx` exposes itself as `docx`. Local import keeps the
    # cold-start path light.
    from docx import Document

    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001
        raise PRDLoaderError(f"Word 文档解析失败：{e}") from e

    # Paragraphs first; then any text inside tables (PRDs often put
    # acceptance criteria in tables).
    lines: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                lines.append(row_text)
    return "\n".join(lines)


def _extract_plain(data: bytes) -> str:
    """Decode plain text with a safe fallback chain."""
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    # latin-1 always succeeds in theory; if we get here something's
    # very wrong and we'd rather not silently return garbage.
    raise PRDLoaderError("文件编码无法识别（尝试过 utf-8 / gb18030 / latin-1）")
